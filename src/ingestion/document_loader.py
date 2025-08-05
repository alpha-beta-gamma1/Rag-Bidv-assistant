import os
import json
from typing import Dict, Any
from pathlib import Path
from docx import Document
import PyPDF2
from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class DocumentLoader:
    """Enhanced document loader with error handling"""
    
    def __init__(self):
        self.supported_formats = {'.docx', '.pdf', '.txt', '.json'}
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load document from file with proper error handling"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            if file_ext == '.docx':
                return self._load_docx(file_path)
            elif file_ext == '.pdf':
                return self._load_pdf(file_path)
            elif file_ext == '.txt':
                return self._load_txt(file_path)
            elif file_ext == '.json':
                return self._load_json(file_path)
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            raise
    
    def _load_docx(self, file_path: str) -> Dict[str, Any]:
        """Load DOCX file"""
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'content': '\n'.join(paragraphs),
            'paragraphs': paragraphs,
            'tables': tables,
            'metadata': {
                'file_path': file_path,
                'file_type': 'docx',
                'paragraphs_count': len(paragraphs),
                'tables_count': len(tables)
            }
        }
    
    def _load_pdf(self, file_path: str) -> Dict[str, Any]:
        """Load PDF file"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
            
            full_text = '\n'.join(text_content)
            
            return {
                'content': full_text,
                'pages': text_content,
                'metadata': {
                    'file_path': file_path,
                    'file_type': 'pdf',
                    'pages_count': len(pdf_reader.pages)
                }
            }
    
    def _load_txt(self, file_path: str) -> Dict[str, Any]:
        """Load text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return {
            'content': content,
            'metadata': {
                'file_path': file_path,
                'file_type': 'txt',
                'length': len(content)
            }
        }
    
    def _load_json(self, file_path: str) -> Dict[str, Any]:
        """Load JSON file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Convert to standard format
        if isinstance(data, list):
            content = json.dumps(data, ensure_ascii=False, indent=2)
        else:
            content = json.dumps(data, ensure_ascii=False, indent=2)
        
        return {
            'content': content,
            'data': data,
            'metadata': {
                'file_path': file_path,
                'file_type': 'json',
                'items_count': len(data) if isinstance(data, list) else 1
            }
        }