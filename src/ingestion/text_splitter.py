import json
import re
from typing import List, Dict, Any
from docx import Document
from pathlib import Path
from src.utils.config import Config
from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class TextSplitter:
    def __init__(self, config: Config):
        self.config = config
        self.max_tokens = config.get('chunking.max_tokens', 400)
        self.overlap = config.get('chunking.overlap', 50)
        
        logger.info(f"TextSplitter initialized with max_tokens={self.max_tokens}, overlap={self.overlap}")

    @staticmethod
    def count_tokens(text: str) -> int:
        """Đếm token cơ bản (ước lượng theo từ)."""
        return len(text.split())

    def split_paragraph(self, text: str) -> List[str]:
        """Tách đoạn dài thành các chunk nhỏ dựa trên số token."""
        sentences = re.split(r'(?<=[.!?]) +', text)
        chunks, current = [], []
        token_count = 0
        
        for sentence in sentences:
            sentence_tokens = self.count_tokens(sentence)
            
            if token_count + sentence_tokens > self.max_tokens and current:
                # Add current chunk
                chunks.append(" ".join(current))
                
                # Start new chunk with overlap if configured
                if self.overlap > 0 and len(current) > 1:
                    # Keep last few sentences for overlap
                    overlap_sentences = []
                    overlap_tokens = 0
                    for s in reversed(current):
                        s_tokens = self.count_tokens(s)
                        if overlap_tokens + s_tokens <= self.overlap:
                            overlap_sentences.insert(0, s)
                            overlap_tokens += s_tokens
                        else:
                            break
                    current = overlap_sentences + [sentence]
                    token_count = overlap_tokens + sentence_tokens
                else:
                    current = [sentence]
                    token_count = sentence_tokens
            else:
                current.append(sentence)
                token_count += sentence_tokens
        
        if current:
            chunks.append(" ".join(current))
        
        return chunks

    @staticmethod
    def parse_table(table, context_title=None) -> Dict[str, Any]:
        """Chuyển bảng trong DOCX thành cấu trúc JSON."""
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        
        return {
            "type": "table",
            "title": context_title if context_title else "Unknown Table",
            "columns": rows[0] if rows else [],
            "rows": rows[1:] if len(rows) > 1 else []
        }
    
    def split_document(self, document: Dict[str, Any]) -> List[Dict[str, Any]]:
        file_path = document.get('file_path', '')
        metadata = document.get('metadata', {})
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.docx':
            logger.info("[INFO] Using DOCX splitter based on file_path")
            return self._split_docx_document(file_path, metadata)
        
        content = document.get('content', '')
        if file_extension == '.json':
            return self._split_json_document(content, metadata)
        else:
            return self._split_text_document(content, metadata)

    def _split_docx_document(self, file_path: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split DOCX document into chunks"""
        try:
            doc = Document(file_path)
            chunks = []
            current_title = None
            buffer = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Phát hiện tiêu đề
                if (re.match(r'^\d+(\.\d+)*\s', text) or 
                    text.isupper() or 
                    para.style.name.startswith("Heading")):
                    
                    # Flush buffer trước đó
                    if buffer:
                        joined_text = " ".join(buffer)
                        for chunk_content in self.split_paragraph(joined_text):
                            chunks.append({
                                "type": "text",
                                "title": current_title or "Untitled",
                                "content": chunk_content,
                                "metadata": {**metadata, "title": current_title}
                            })
                        buffer = []
                    current_title = text
                else:
                    buffer.append(text)

            # Flush cuối cùng
            if buffer:
                joined_text = " ".join(buffer)
                for chunk_content in self.split_paragraph(joined_text):
                    chunks.append({
                        "type": "text",
                        "title": current_title or "Untitled",
                        "content": chunk_content,
                        "metadata": {**metadata, "title": current_title}
                    })

            # Xử lý bảng
            for table in doc.tables:
                table_chunk = self.parse_table(table, context_title=current_title)
                table_chunk["metadata"] = {**metadata, "title": current_title}
                chunks.append(table_chunk)

            logger.info(f"Split DOCX document into {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error splitting DOCX document: {e}")
            raise

    def _split_json_document(self, content: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split JSON document (already chunked)"""
        try:
            chunks_data = json.loads(content)
            
            # Handle different JSON structures
            if isinstance(chunks_data, dict):
                for key in ['chunks', 'data', 'items']:
                    if key in chunks_data and isinstance(chunks_data[key], list):
                        chunks_data = chunks_data[key]
                        break
            
            if not isinstance(chunks_data, list):
                raise ValueError("JSON must contain a list of chunks")
            
            processed_chunks = []
            for i, chunk in enumerate(chunks_data):
                if isinstance(chunk, dict):
                    # Add metadata to each chunk
                    chunk_with_metadata = chunk.copy()
                    chunk_with_metadata["metadata"] = {
                        **metadata, 
                        "chunk_id": i,
                        "original_title": chunk.get("title", "")
                    }
                    processed_chunks.append(chunk_with_metadata)
            
            logger.info(f"Processed JSON with {len(processed_chunks)} chunks")
            return processed_chunks

        except Exception as e:
            logger.error(f"Error processing JSON document: {e}")
            raise

    def _split_text_document(self, content: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split plain text document"""
        if not content.strip():
            return []
        
        chunks = []
        text_chunks = self.split_paragraph(content)
        
        for i, chunk_content in enumerate(text_chunks):
            chunks.append({
                "type": "text",
                "title": f"Text Chunk {i+1}",
                "content": chunk_content,
                "metadata": {**metadata, "chunk_id": i}
            })
        
        logger.info(f"Split text document into {len(chunks)} chunks")
        return chunks